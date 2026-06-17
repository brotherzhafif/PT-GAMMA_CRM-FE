from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\yudha\Downloads\GeoCoast AI_ringkas.docx")


def set_spacing(paragraph, before=0, after=6, line=1.10):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    return run


def add_para(doc, text, style=None, after=6, justify=False):
    p = doc.add_paragraph(style=style)
    add_run(p, text)
    set_spacing(p, after=after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, item)
        set_spacing(p, after=4, line=1.167)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_spacing(p, before=16 if level == 1 else 12, after=8 if level == 1 else 6)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

list_style = styles["List Bullet"]
list_style.font.name = "Calibri"
list_style.font.size = Pt(11)
list_style.paragraph_format.space_after = Pt(8)
list_style.paragraph_format.line_spacing = 1.167

title = doc.add_paragraph()
set_spacing(title, after=8, line=1.10)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("GeoCoast AI: Sistem Cerdas untuk Prioritisasi Risiko, Estimasi Kerugian, dan Rekomendasi Mitigasi Abrasi Pesisir Berbasis Data Geospasial")
r.bold = True
r.font.name = "Calibri"
r.font.size = Pt(15)
r.font.color.rgb = RGBColor.from_string("0B2545")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(meta, after=10)
add_run(meta, "Universitas Ahmad Dahlan Yogyakarta | 2026")

add_heading(doc, "ABSTRACT")
add_para(
    doc,
    "Abrasi pantai mengancam keberlanjutan wilayah pesisir Indonesia melalui kehilangan daratan, kerusakan infrastruktur, penurunan aktivitas ekonomi, dan degradasi ekosistem. Pemanfaatan penginderaan jauh dan kecerdasan buatan telah meningkatkan kemampuan pemantauan perubahan garis pantai, tetapi hasilnya masih terbatas untuk mendukung prioritisasi kebijakan mitigasi. Penelitian ini mengusulkan GeoCoast AI, sistem pendukung keputusan yang mengintegrasikan data geospasial, karakteristik pesisir, estimasi kerugian, dan basis pengetahuan mitigasi untuk menghasilkan prioritas risiko abrasi, rekomendasi tindakan, serta penjelasan ilmiah yang dapat ditelusuri. Analisis dilakukan melalui indeks risiko berbasis multikriteria, estimasi potensi kerugian, dan rekomendasi berbasis Retrieval-Augmented Generation (RAG). Luaran penelitian berupa prototipe sistem yang membantu pemerintah dan pemangku kepentingan menetapkan prioritas penanganan abrasi secara objektif, efisien, dan transparan.",
    justify=True,
)
add_para(doc, "Kata kunci: abrasi pantai, geospasial, sistem pendukung keputusan, estimasi kerugian, explainable AI.", after=8)

add_heading(doc, "PENDAHULUAN")
add_para(
    doc,
    "Indonesia memiliki garis pantai lebih dari 108.000 kilometer dan kawasan pesisir yang menopang permukiman, pariwisata, pelabuhan, perikanan, infrastruktur publik, serta ekosistem bernilai tinggi. Namun, kawasan tersebut semakin rentan terhadap abrasi akibat dinamika pantai, kenaikan muka air laut, perubahan iklim, dan tekanan pembangunan. Dampaknya mencakup kehilangan lahan, kerusakan infrastruktur, gangguan mata pencaharian masyarakat pesisir, serta peningkatan risiko sosial-ekonomi dan ekologis.",
    justify=True,
)
add_para(
    doc,
    "Perkembangan teknologi geospasial dan Artificial Intelligence (AI) telah memperkuat pemantauan abrasi. BRIN (2026), misalnya, mengembangkan metode AI untuk pemetaan garis pantai utara Jawa dengan akurasi dan kedetailan lebih baik dibandingkan pendekatan konvensional. Meski demikian, sebagian besar kajian masih berfokus pada deteksi dan monitoring perubahan garis pantai, sementara kebutuhan pengambil kebijakan adalah informasi yang lebih operasional: wilayah mana yang perlu diprioritaskan, berapa potensi kerugiannya, dan strategi mitigasi apa yang paling sesuai.",
    justify=True,
)
add_para(
    doc,
    "Kesenjangan tersebut penting karena sumber daya dan anggaran mitigasi terbatas. Tanpa sistem prioritisasi yang objektif, penanganan abrasi berisiko tidak tepat sasaran, sulit dipertanggungjawabkan, dan kurang adaptif terhadap karakteristik wilayah. Kajian sistem pendukung keputusan pesisir menunjukkan perlunya integrasi indeks, data spasial, dan instrumen rekomendasi untuk perencanaan pesisir yang berkelanjutan (Barzehkar et al., 2021). Oleh karena itu, penelitian ini menempatkan novelty pada penggabungan prioritisasi risiko abrasi, estimasi kerugian, rekomendasi mitigasi, dan explainable AI dalam satu prototipe GeoCoast AI.",
    justify=True,
)

add_heading(doc, "TUJUAN")
add_para(
    doc,
    "Tujuan umum penelitian ini adalah mengembangkan GeoCoast AI sebagai sistem pendukung keputusan berbasis kecerdasan buatan untuk memprioritaskan risiko abrasi pesisir, mengestimasi potensi kerugian, dan merekomendasikan mitigasi yang sesuai dengan karakteristik wilayah secara transparan dan dapat dijelaskan secara ilmiah.",
    justify=True,
)
add_bullets(
    doc,
    [
        "Mengidentifikasi risiko abrasi berdasarkan data geospasial, perubahan garis pantai, dan karakteristik fisik pesisir.",
        "Menganalisis kerentanan wilayah dengan mempertimbangkan penggunaan lahan, kepadatan penduduk, infrastruktur, dan fungsi kawasan.",
        "Mengestimasi potensi kerugian sebagai dasar prioritas investasi mitigasi.",
        "Menghasilkan rekomendasi mitigasi berbasis AI dan knowledge retrieval yang disertai penjelasan ilmiah.",
    ],
)

add_heading(doc, "DESKRIPSI RISET")
add_para(
    doc,
    "GeoCoast AI dirancang sebagai prototipe sistem pendukung keputusan yang mengolah data perubahan garis pantai, karakteristik fisik pesisir, penggunaan lahan, kepadatan penduduk, infrastruktur strategis, dan fungsi kawasan seperti wisata, pelabuhan, permukiman, tambak, serta konservasi. Data tersebut diintegrasikan dalam basis data geospasial untuk menghitung indeks prioritas risiko pada setiap unit wilayah pesisir.",
    justify=True,
)
add_para(
    doc,
    "Sistem tidak berhenti pada pemetaan lokasi terdampak. GeoCoast AI menambahkan estimasi potensi kerugian berbasis luas terdampak, nilai ekonomi penggunaan lahan, infrastruktur, dan aktivitas masyarakat pesisir. Setelah prioritas ditentukan, sistem menggunakan basis pengetahuan yang berasal dari literatur ilmiah, regulasi, pedoman pengelolaan pesisir, dan studi kasus mitigasi untuk menghasilkan rekomendasi tindakan. Pendekatan RAG memungkinkan rekomendasi dikaitkan dengan bukti yang relevan sehingga pengguna memahami alasan pemilihan strategi, bukan hanya menerima keluaran sistem.",
    justify=True,
)

add_heading(doc, "METODE")
add_para(
    doc,
    "Penelitian menggunakan pendekatan Research and Development (R&D) untuk membangun dan mengevaluasi prototipe GeoCoast AI. Tahap pertama adalah pengumpulan dan integrasi data geospasial, meliputi perubahan garis pantai, penggunaan lahan, kepadatan penduduk, infrastruktur, fungsi kawasan, karakteristik fisik pantai, serta literatur mitigasi abrasi.",
    justify=True,
)
add_para(
    doc,
    "Tahap kedua adalah analisis risiko abrasi melalui parameter perubahan garis pantai, karakteristik fisik pesisir, kerentanan penggunaan lahan, kepadatan penduduk, dan keberadaan infrastruktur penting. Parameter tersebut sejalan dengan pendekatan Coastal Vulnerability Index (CVI) dalam kajian kerentanan pesisir berbasis penginderaan jauh dan GIS (Hastuti & Nagai, 2022). Pembobotan parameter direncanakan menggunakan Analytical Hierarchy Process (AHP) karena mampu mengakomodasi penilaian multikriteria secara sistematis dan telah digunakan dalam analisis kerentanan pesisir (Nanda et al., 2024).",
    justify=True,
)
add_para(
    doc,
    "Tahap ketiga adalah estimasi potensi kerugian berdasarkan luas terdampak, nilai ekonomi penggunaan lahan, aset infrastruktur, dan aktivitas ekonomi masyarakat. Tahap keempat adalah penyusunan rekomendasi mitigasi menggunakan RAG, yaitu penggabungan Large Language Model dengan basis pengetahuan mitigasi abrasi. Tahap kelima adalah evaluasi sistem melalui kesesuaian rekomendasi terhadap karakteristik wilayah, konsistensi keluaran, kemudahan interpretasi, dan validasi oleh ahli atau literatur yang relevan.",
    justify=True,
)

add_heading(doc, "POTENSI DAMPAK")
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Aspek"
hdr[1].text = "Dampak yang Diharapkan"
for cell in hdr:
    shade_cell(cell, "F2F4F7")
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
rows = [
    ("Ilmu pengetahuan", "Memperkuat kajian sistem pendukung keputusan pesisir melalui integrasi analisis risiko, estimasi kerugian, knowledge retrieval, dan explainable AI."),
    ("Teknologi", "Menghasilkan prototipe yang memanfaatkan GIS, AI, dan RAG untuk mengubah data geospasial menjadi prioritas serta rekomendasi mitigasi yang dapat ditelusuri."),
    ("Lingkungan", "Mendukung pemilihan strategi perlindungan pantai yang lebih sesuai dengan karakteristik wilayah sehingga risiko kerusakan ekosistem pesisir dapat dikurangi."),
    ("Ekonomi dan kebijakan", "Membantu alokasi anggaran mitigasi berdasarkan risiko dan potensi kerugian, sekaligus memperkuat transparansi, objektivitas, dan akuntabilitas keputusan publik."),
]
for a, b in rows:
    cells = table.add_row().cells
    cells[0].text = a
    cells[1].text = b
set_table_width(table, [2300, 7060])
set_cell_margins(table)
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            set_spacing(p, after=2, line=1.10)

add_heading(doc, "KESIMPULAN")
add_para(
    doc,
    "GeoCoast AI menjawab keterbatasan penelitian abrasi yang selama ini dominan pada deteksi dan monitoring dengan menawarkan sistem pendukung keputusan yang lebih operasional. Melalui integrasi data geospasial, indeks risiko multikriteria, estimasi potensi kerugian, dan rekomendasi mitigasi berbasis RAG, sistem ini diharapkan mampu membantu pengambil kebijakan menentukan wilayah prioritas dan strategi penanganan yang efektif. Keunggulan utama penelitian terletak pada kemampuan menghasilkan rekomendasi yang tidak hanya relevan secara teknis, tetapi juga dapat dijelaskan secara ilmiah sehingga mendukung pengelolaan pesisir yang objektif, transparan, dan berkelanjutan.",
    justify=True,
)

p = doc.add_paragraph()
p.add_run().add_break(WD_BREAK.PAGE)
add_heading(doc, "DAFTAR PUSTAKA")
refs = [
    "Barzehkar, M., Parnell, K., Soomere, T., Dragovich, D., & Engström, J. (2021). Decision support tools, systems and indices for sustainable coastal planning and management: A review. Ocean & Coastal Management, 212, 105813. https://doi.org/10.1016/j.ocecoaman.2021.105813",
    "BRIN, B. R. I. N. (2026). BRIN Kembangkan Metode AI untuk Pemetaan Garis Pantai Utara Jawa Lebih Akurat dan Detail. 1 April 2026. https://brin.go.id/press-release/127271/brin-kembangkan-metode-ai-untuk-pemetaan-garis-pantai-utara-jawa-lebih-akurat-dan-detail",
    "Hastuti, A. W., & Nagai, M. (2022). Coastal Vulnerability Assessment of Bali Province, Indonesia Using Remote Sensing and GIS Approaches.",
    "Nanda, G., Ulfa, A., Adriadi, R., & Anjas, H. (2024). Coastal Vulnerability Assessment Based on Coastal Vulnerability Index (CVI) on the Coastal Area of Kolaka Regency, Southeast Sulawesi, Indonesia. Cvi, 267-279.",
]
for ref in refs:
    p = add_para(doc, ref, after=4)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.left_indent = Inches(0.25)

doc.save(OUT)
print(OUT)
